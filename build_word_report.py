# -*- coding: utf-8 -*-
"""
==============================================================================
 SCRIPT DE GERAÇÃO DO RELATÓRIO SIMPLIFICADO EM WORD (.DOCX)
 ATIVIDADE DE EMBEDDED C (VÍDEOS 1 E 2) — UFU / FEELT
==============================================================================
 Discente: Vicenzo De Marco Olivalves (12421ECP006)
 Disciplina: Sistemas Embarcados I — Prof. Jeovane Vicente de Sousa
"""

import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def create_simplified_word_document():
    doc = docx.Document()

    # Margens Padrão (2.5 cm)
    for section in doc.sections:
        section.top_margin = Inches(0.98)
        section.bottom_margin = Inches(0.98)
        section.left_margin = Inches(0.98)
        section.right_margin = Inches(0.98)

    # Cores
    c_primary = RGBColor(35, 38, 121)    # Azul #232679
    c_dark = RGBColor(15, 20, 60)       # Escuro
    c_body = RGBColor(31, 41, 55)       # Texto normal #1f2937

    # TÍTULO E CABEÇALHO
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("UNIVERSIDADE FEDERAL DE UBERLÂNDIA\nFACULDADE DE ENGENHARIA ELÉTRICA — FEELT\nDISCIPLINA DE SISTEMAS EMBARCADOS I")
    r_title.bold = True
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(11)
    r_title.font.color.rgb = c_primary

    p_main_title = doc.add_paragraph()
    p_main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_main_title.paragraph_format.space_before = Pt(16)
    p_main_title.paragraph_format.space_after = Pt(12)
    r_main = p_main_title.add_run("RELATÓRIO TÉCNICO DE ESTUDO — EMBEDDED C\n(VÍDEOS 1 E 2)")
    r_main.bold = True
    r_main.font.name = 'Calibri'
    r_main.font.size = Pt(15)
    r_main.font.color.rgb = c_primary

    # IDENTIFICAÇÃO DO ALUNO E PROFESSOR
    table_meta = doc.add_table(rows=1, cols=2)
    table_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_l, cell_r = table_meta.rows[0].cells
    cell_l.width = Inches(3.2)
    cell_r.width = Inches(3.2)

    p_l = cell_l.paragraphs[0]
    p_l.add_run("Discente: ").bold = True
    p_l.add_run("Vicenzo De Marco Olivalves\n")
    p_l.add_run("Matrícula: ").bold = True
    p_l.add_run("12421ECP006\n")
    p_l.add_run("Curso: ").bold = True
    p_l.add_run("Engenharia de Computação")
    for r in p_l.runs:
        r.font.name = 'Calibri'
        r.font.size = Pt(10)

    p_r = cell_r.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_r.add_run("Docente Responsável:\n").bold = True
    p_r.add_run("Prof. Jeovane Vicente de Sousa\n")
    p_r.add_run("Semestre: ").bold = True
    p_r.add_run("2026/1")
    for r in p_r.runs:
        r.font.name = 'Calibri'
        r.font.size = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.bold = True
        r.font.name = 'Calibri'
        r.font.size = Pt(13)
        r.font.color.rgb = c_primary

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.bold = True
        r.font.name = 'Calibri'
        r.font.size = Pt(11)
        r.font.color.rgb = c_dark

    def add_body(text, bold_prefix=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_pre = p.add_run(bold_prefix)
            r_pre.bold = True
            r_pre.font.name = 'Calibri'
            r_pre.font.size = Pt(11)
            r_pre.font.color.rgb = c_body
        r = p.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(11)
        r.font.color.rgb = c_body

    def add_figure(img_path, caption_text):
        if os.path.exists(img_path):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(8)
            p_img.paragraph_format.space_after = Pt(2)
            run = p_img.add_run()
            run.add_picture(img_path, width=Inches(5.0))
            
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_after = Pt(10)
            r_cap = p_cap.add_run(caption_text)
            r_cap.italic = True
            r_cap.font.name = 'Calibri'
            r_cap.font.size = Pt(9.5)
            r_cap.font.color.rgb = RGBColor(100, 100, 100)

    # --------------------------------------------------------------------------
    # DESENVOLVIMENTO DO RELATÓRIO (VERSÃO SIMPLIFICADA)
    # --------------------------------------------------------------------------

    # SEÇÃO 1 - VÍDEO 1: CONCEITO E ESTRUTURA
    add_h1("1. Análise do Vídeo 1: Introdução ao Embedded C e Arquitetura")
    
    # Texto manuscrito do aluno Vicenzo
    add_body("No primeiro vídeo, o professor Marcelo começa explicando o porquê aprender C para sistemas embarcados exige mudar a forma de pensar dos programadores, apresentando a diferença essencial entre escrever um programa em C de forma tradicional no computador e escrever um programa em C em um microcontrolador, como no caso de uma placa STM32F103 Blue Pill. No caso dessas diferenças, é que no desktop o programa roda em cima de um sistema operacional, como Windows, Linux e afins. A linguagem C é usada para criar aplicativos de alto nível. O sistema operacional aloca a memória RAM virtualizada para o seu programa, isola o processo e gerencia os arquivos. Se o programa tentar acessar um endereço de memória que não é dele, o sistema operacional intervém e dispara um erro de segmentação.")

    add_body("Agora no caso de um sistema embarcado, o código roda direto na placa, então não tem essa rede de proteção, pois não tem a presença de um sistema operacional. O programador tem total controle e responsabilidade sobre o hardware e afins. É o programador que configura o clock, barramentos, direção de pinos de saída e entrada (GPIO), interrupções e afins, igual estamos fazendo na matéria.")

    add_body("Junto dessa ideia de rodar o código diretamente no hardware sem um sistema operacional por cima, a aula avança para a estrutura do fluxo de execução na função main. Em um código C comum, a função \"int main(void)\" é chamada pelo sistema operacional, executa uma sequência e volta retornando 0 e a CPU volta a outras funções. No caso de um microcontrolador embarcado, a função main() não pode chegar ao fim, pois se chegar a CPU entra em um estado indefinido, o ponteiro cai para um endereço inválido e o sistema trava ou reinicia. Por conta disso a arquitetura do firmware tem duas etapas obrigatórias:")

    add_body("Roda apenas uma vez quando a placa é ligada ou resetada e é onde configuramos o clock, inicializamos os pinos de GPIO como entrada ou saída e ligamos os periféricos.", bold_prefix="• Setup: ")
    add_body("Como resposta àquela ideia de que não pode sair da main, após o código inicializar, ele fica em um laço de while(1), aí o processador fica executando tarefas continuamente ou entra em modo de repouso esperando por uma interrupção.", bold_prefix="• Loop: ")

    # FIGURA 1 - LINHA DO TEMPO C
    add_figure(r'C:\Users\vicen\Downloads\Video\frames_v1\frame_03m59s.png', "Figura 1: Slide exibido no minuto 03:59 apresentando a linha do tempo e evolução dos padrões da linguagem C.")

    add_h2("1.1 Histórico dos Padrões C e a Biblioteca <stdint.h>")
    add_body("Na continuidade do vídeo, o professor mostra a evolução dos padrões da linguagem C (K&R 1978, C89/C90, C99, C11, C17). O ponto de maior destaque para sistemas embarcados é o surgimento do padrão C99 (1999), que introduziu a biblioteca <stdint.h>. Ela permitiu criar tipos de dados de tamanho fixo garantido (uint8_t, uint16_t, uint32_t), evitando o uso de int genéricos e garantindo que o tipo corresponda exatamente à largura de 32 bits dos registradores da família ARM Cortex-M.")

    add_h2("1.2 Organização da Memória no Microcontrolador")
    add_body("A aula detalha como o compilador distribui o programa na memória física do STM32:")
    add_body("Armazena permanentemente o binário do código e as constantes (const). Não apaga ao desligar a placa.", bold_prefix="• Memória Flash ROM (.text): ")
    add_body("Guarda as variáveis globais e estáticas inicializadas. É carregada na inicialização.", bold_prefix="• Memória SRAM (.data): ")
    add_body("Guarda as variáveis globais não inicializadas, que são zeradas durante o boot.", bold_prefix="• Memória SRAM (.bss): ")
    add_body("Armazena variáveis locais e registradores temporários de funções. O professor faz um alerta sobre o risco de estouro da pilha (Stack Overflow) quando se abusa de variáveis locais grandes ou recursão.", bold_prefix="• Pilha (Stack): ")

    add_figure(r'C:\Users\vicen\Downloads\Video\frames_v1\frame_61m59s.png', "Figura 2: Slide do minuto 61:59 detalhando as seções de memória (.text, .data, .bss) no microcontrolador.")

    add_h2("1.3 Qualificadores Especiais: volatile, const e static")
    add_body("O professor dá um grande destaque ao qualificador volatile. Ele avisa ao compilador que uma variável pode ser alterada fora do fluxo normal do código (por uma interrupção de hardware ISR ou alteração física em um pino). O volatile impede que o compilador faça otimizações incorretas salvando o valor em um registrador interno da CPU. O qualificador const salva tabelas na Flash ROM para economizar RAM, enquanto o static preserva o valor de variáveis locais entre chamadas de função e atua como encapsulamento privado dentro do arquivo .c.")

    add_h2("1.4 Manipulação de Registradores por Operações Bitwise")
    add_body("Em microcontroladores, a alteração de pinos é feita manipulando registradores de 32 bits através de operações bitwise sem destruir as configurações dos pinos vizinhos: usa-se OR (|) para ligar um pino (SET), AND com NOT (& ~) para desligar um pino (CLEAR), XOR (^) para inverter o estado do pino (TOGGLE), e MASK com AND (&) para testar se um pino de entrada está ativo.")

    # --------------------------------------------------------------------------
    # SEÇÃO 2 - VÍDEO 2: ABSTRAÇÃO E ESTRUTURAS AVANÇADAS
    # --------------------------------------------------------------------------
    add_h1("2. Análise do Vídeo 2: Mapeamento de Memória, Estruturas e Interrupções")
    
    add_body("Na segunda videoaula, o professor avança da manipulação de registradores para o uso de abstrações profissionais com Estruturas (struct), Unióes (union), Enumerações (enum) e Ponteiros de Função.")

    add_h2("2.1 Mapeamento de Registradores com struct (Memory-Mapped I/O)")
    add_body("O mapa de memória do ARM Cortex-M coloca os registradores de um periférico (como a porta GPIOA) em posições contíguas sequenciais (offsets 0x00, 0x04, 0x08, 0x0C). Ao criar uma struct que espelha essa sequência, podemos apontar um ponteiro para o endereço base da porta (0x40010800) e acessar os pinos com uma sintaxe limpa (GPIOA->ODR), exatamente como é feito nas bibliotecas ST HAL e CMSIS:")

    add_figure(r'C:\Users\vicen\Downloads\Video\frames_v2\frame_70m59s.png', "Figura 3: Slide do minuto 70:59 demonstrando a organização de dados e periféricos através de estruturas (struct).")

    add_h2("2.2 Compartilhamento de Memória com union")
    add_body("Diferente da struct, na union todos os membros compartilham a mesma posição física na memória RAM. Isso é muito útil para fatiar um dado numérico de 32 bits (uint32_t) em 4 bytes separados (uint8_t) para envio por pacotes USB HID ou comunicação serial UART sem precisar de máscaras manuais de bitwise:")

    add_figure(r'C:\Users\vicen\Downloads\Video\frames_v2\frame_78m59s.png', "Figura 4: Slide do minuto 78:59 explicando o funcionamento de Unióes (union) e o compartilhamento de RAM.")

    add_h2("2.3 Máquinas de Estados Finitos (FSM) com enum")
    add_body("Substituir números inteiros soltos por nomes de estados legíveis (enum) permite construir Máquinas de Estados Finitos (FSM) conectadas a laços switch-case, aumentando a segurança do código e evitando que o sistema caia em estados inválidos:")

    add_figure(r'C:\Users\vicen\Downloads\Video\frames_v2\frame_74m59s.png', "Figura 5: Slide do minuto 74:59 demonstrando a organização de enumerações (enum) para controle de fluxo.")

    add_h2("2.4 Ponteiros para Função, Tabela IVT e Callbacks da HAL")
    add_body("O nome de uma função em C nada mais é do que um ponteiro para o endereço de memória Flash onde o código executável começa. Isso se relaciona com a Tabela de Vetores de Interrupção (IVT alocada no endereço 0x08000000 da Flash): quando ocorre uma interrupção de hardware (como o estouro do Timer TIM2), a CPU lê o ponteiro da tabela e salta direto para a função ISR de callback (ex.: HAL_TIM_PeriodElapsedCallback).")

    add_figure(r'C:\Users\vicen\Downloads\Video\frames_v2\frame_38m59s.png', "Figura 6: Slide do minuto 38:59 apresentando a sintaxe e aplicação de Ponteiros para Função.")

    add_h2("2.5 Demonstração Prática no Godbolt Compiler Explorer")
    add_body("No minuto 62:59 do vídeo 2, o professor abre ao vivo o site Godbolt Compiler Explorer (godbolt.org) com o compilador ARM GCC 11.2 (none). Ele mostra na prática o código Assembly gerado pelo compilador para demonstrar como o uso dos modificadores volatile e const altera diretamente as instruções de carregamento (LDR/STR) geradas para o processador ARM Cortex-M.")

    add_figure(r'C:\Users\vicen\Downloads\Video\frames_v2\frame_62m59s.png', "Figura 7: Demonstração ao vivo no Godbolt Compiler Explorer (minuto 62:59) analisando o código Assembly ARM GCC.")

    # --------------------------------------------------------------------------
    # SEÇÃO 3 - CONCLUSÃO E SÍNTESE
    # --------------------------------------------------------------------------
    add_h1("3. Conclusão e Síntese Comparativa")
    add_body("Assistir às duas videoaulas permitiu entender perfeitamente como a programação em C para sistemas embarcados exige um rigor muito maior do que a programação tradicional para computadores. Em embarcados, o programador lida diretamente com o silício, controlando o tempo de execução por interrupções, otimizando o uso de memória Flash e SRAM, e manipulando registradores de hardware de forma segura.")

    add_body("Abaixo apresento a tabela resumo comparando as principais diferenças observadas entre os dois ambientes:")

    # TABELA COMPARATIVA FINAL
    table_comp = doc.add_table(rows=1, cols=3)
    table_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table_comp.rows[0].cells
    hdr_cells[0].text = "Critério Técnico"
    hdr_cells[1].text = "C para PC (Desktop)"
    hdr_cells[2].text = "Embedded C (Embarcados)"

    for c in hdr_cells:
        set_cell_background(c, "232679")
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = 'Calibri'
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

    data = [
        ("Ambiente de Execução", "Aplicativo sobre Sistema Operacional", "Firmware Bare-Metal no silício"),
        ("Acesso ao Hardware", "Bloqueado pelo SO / Via Syscalls", "Direto por registradores MMIO"),
        ("Ciclo de Vida (main)", "Retorna 0 para o SO ao terminar", "Loop infinito while(1) ou repouso __WFI()"),
        ("Determinismo de Tempo", "Não-determinístico (escalonado pelo SO)", "Tempo Real determinístico por ISR/NVIC"),
        ("Qualificador volatile", "Raramente necessário", "Obrigatório em ISRs e registradores"),
        ("Tipos de Dados Inteiros", "Uso de int de tamanho variável", "Obrigatório <stdint.h> (uint8_t, uint32_t)"),
        ("Gerenciamento de RAM", "Abundante com memória virtual", "Estrito e contado (SRAM vs Flash ROM)")
    ]

    for row_data in data:
        row = table_comp.add_row()
        for idx, text in enumerate(row_data):
            cell = row.cells[idx]
            cell.text = text
            if idx == 0:
                set_cell_background(cell, "F1F5F9")
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx != 0 else WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = 'Calibri'
                    r.font.size = Pt(9.5)
                    if idx == 0:
                        r.bold = True

    # SALVAR ARQUIVO DOCX
    out_docx1 = r'C:\Users\vicen\Downloads\Vicenzo_Olivalves_Relatorio_Embedded_C_UFU.docx'
    out_docx2 = r'c:\Users\vicen\Downloads\streamdeck\Vicenzo_Olivalves_Relatorio_Embedded_C_UFU.docx'
    doc.save(out_docx1)
    doc.save(out_docx2)
    print(f"Successfully created SIMPLIFIED Word Report: {out_docx1}")

if __name__ == '__main__':
    create_simplified_word_document()
