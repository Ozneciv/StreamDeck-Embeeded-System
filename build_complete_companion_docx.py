# -*- coding: utf-8 -*-
"""
==============================================================================
 SCRIPT DE GERAÇÃO DO RELATÓRIO EXAUSTIVO COM PONTO FLUTUANTE (.DOCX)
 ATIVIDADE DE EMBEDDED C (VÍDEOS 1 E 2) — UFU / FEELT
==============================================================================
 Discente: Vicenzo De Marco Olivalves (12421ECP006)
 Disciplina: Sistemas Embarcados I — Prof. Jeovane Vicente de Sousa
"""

import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def create_exhaustive_word_document_v2():
    doc = docx.Document()

    # Margens (2.5 cm)
    for section in doc.sections:
        section.top_margin = Inches(0.98)
        section.bottom_margin = Inches(0.98)
        section.left_margin = Inches(0.98)
        section.right_margin = Inches(0.98)

    # Cores
    c_primary = RGBColor(35, 38, 121)    # Azul #232679
    c_dark = RGBColor(15, 20, 60)       # Escuro
    c_body = RGBColor(31, 41, 55)       # Texto normal #1f2937

    # TÍTULO E CABEÇALHO
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("UNIVERSIDADE FEDERAL DE UBERLÂNDIA\nFACULDADE DE ENGENHARIA ELÉTRICA — FEELT\nDISCIPLINA DE SISTEMAS EMBARCADOS I")
    r_title.bold = True
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(11)
    r_title.font.color.rgb = c_primary

    p_main_title = doc.add_paragraph()
    p_main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_main_title.paragraph_format.space_before = Pt(16)
    p_main_title.paragraph_format.space_after = Pt(12)
    r_main = p_main_title.add_run("RELATÓRIO TÉCNICO E GUIA COMPLETO DE ESTUDO\nEMBEDDED C (VÍDEOS 1 E 2)")
    r_main.bold = True
    r_main.font.name = 'Calibri'
    r_main.font.size = Pt(15)
    r_main.font.color.rgb = c_primary

    # IDENTIFICAÇÃO DO ALUNO E PROFESSOR
    table_meta = doc.add_table(rows=1, cols=2)
    table_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_l, cell_r = table_meta.rows[0].cells
    cell_l.width = Inches(3.2)
    cell_r.width = Inches(3.2)

    p_l = cell_l.paragraphs[0]
    p_l.add_run("Discente: ").bold = True
    p_l.add_run("Vicenzo De Marco Olivalves\n")
    p_l.add_run("Matrícula: ").bold = True
    p_l.add_run("12421ECP006\n")
    p_l.add_run("Curso: ").bold = True
    p_l.add_run("Engenharia de Computação")
    for r in p_l.runs:
        r.font.name = 'Calibri'
        r.font.size = Pt(10)

    p_r = cell_r.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_r.add_run("Docente Responsável:\n").bold = True
    p_r.add_run("Prof. Jeovane Vicente de Sousa\n")
    p_r.add_run("Semestre: ").bold = True
    p_r.add_run("2026/1")
    for r in p_r.runs:
        r.font.name = 'Calibri'
        r.font.size = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.bold = True
        r.font.name = 'Calibri'
        r.font.size = Pt(13)
        r.font.color.rgb = c_primary

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.bold = True
        r.font.name = 'Calibri'
        r.font.size = Pt(11)
        r.font.color.rgb = c_dark

    def add_body(text, bold_prefix=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_pre = p.add_run(bold_prefix)
            r_pre.bold = True
            r_pre.font.name = 'Calibri'
            r_pre.font.size = Pt(11)
            r_pre.font.color.rgb = c_body
        r = p.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(11)
        r.font.color.rgb = c_body

    def add_figure(img_path, caption_text):
        if os.path.exists(img_path):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(8)
            p_img.paragraph_format.space_after = Pt(2)
            run = p_img.add_run()
            run.add_picture(img_path, width=Inches(5.0))
            
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_after = Pt(10)
            r_cap = p_cap.add_run(caption_text)
            r_cap.italic = True
            r_cap.font.name = 'Calibri'
            r_cap.font.size = Pt(9.5)
            r_cap.font.color.rgb = RGBColor(100, 100, 100)

    # --------------------------------------------------------------------------
    # VÍDEO 1: CONCEITOS E ARQUITETURA
    # --------------------------------------------------------------------------
    add_h1("1. Acompanhamento Detalhado do Vídeo 1: Embedded C - Parte 1/2")

    # Texto manuscrito do aluno Vicenzo
    add_body("No primeiro vídeo, o professor Marcelo começa explicando o porquê aprender C para sistemas embarcados exige mudar a forma de pensar dos programadores, apresentando a diferença essencial entre escrever um programa em C de forma tradicional no computador e escrever um programa em C em um microcontrolador, como no caso de uma placa STM32F103 Blue Pill. No caso dessas diferenças, é que no desktop o programa roda em cima de um sistema operacional, como Windows, Linux e afins. A linguagem C é usada para criar aplicativos de alto nível. O sistema operacional aloca a memória RAM virtualizada para o seu programa, isola o processo e gerencia os arquivos. Se o programa tentar acessar um endereço de memória que não é dele, o sistema operacional intervém e dispara um erro de segmentação.")

    add_body("Agora no caso de um sistema embarcado, o código roda direto na placa, então não tem essa rede de proteção, pois não tem a presença de um sistema operacional. O programador tem total controle e responsabilidade sobre o hardware e afins. É o programador que configura o clock, barramentos, direção de pinos de saída e entrada (GPIO), interrupções e afins, igual estamos fazendo na matéria.")

    add_body("Junto dessa ideia de rodar o código diretamente no hardware sem um sistema operacional por cima, a aula avança para a estrutura do fluxo de execução na função main. Em um código C comum, a função \"int main(void)\" é chamada pelo sistema operacional, executa uma sequência e volta retornando 0 e a CPU volta a outras funções. No caso de um microcontrolador embarcado, a função main() não me pode chegar ao fim, pois se chegar a CPU entra em um estado indefinido, o ponteiro cai para um endereço inválido e o sistema trava ou reinicia. Por conta disso a arquitetura do firmware tem duas etapas obrigatórias:")

    add_body("Roda apenas uma vez quando a placa é ligada ou resetada e é onde configuramos o clock, inicializamos os pinos de GPIO como entrada ou saída e ligamos os periféricos.", bold_prefix="• Setup: ")
    add_body("Como resposta àquela ideia de que não pode sair da main, após o código inicializar, ele fica em um laço de while(1), aí o processador fica executando tarefas continuamente ou entra em modo de repouso esperando por uma interrupção.", bold_prefix="• Loop: ")

    add_figure(r'C:\Users\vicen\Downloads\Video\frames_v1\frame_03m59s.png', "Figura 1: Slide do minuto 03:59 apresentando a linha do tempo e a evolução histórica dos padrões da linguagem C.")

    add_h2("1.1 Ponto Flutuante (float / double) vs. Aritmética de Ponto Fixo (Sem FPU)")
    add_body("Logo no início do vídeo, o professor faz um comentário crucial sobre o uso de Ponto Flutuante (float e double) em microcontroladores. Ele explica que microcontroladores mais simples ou de custo reduzido (como o ARM Cortex-M3 do STM32F103) **não possuem uma FPU de hardware (Floating Point Unit)**. Quando o programador escreve operações matematicas usando float ou double no C, o compilador é forçado a incluir uma biblioteca de software (Soft-Float) gigantesca dentro da memória Flash ROM para emular as operações em ponto flutuante por instrução de software. Isso resulta em um consumo enorme de memória Flash e faz com que uma simples multiplicação leve dezenas de ciclos de clock extras de CPU. Por isso, a norma de Embedded C (ISO/IEC TR 18037) preconiza o uso de Aritmética de Ponto Fixo (Fixed-Point Arithmetic) ou o uso de inteiros escalonados (ex.: medir milivolt em vez de volt como float) para garantir máxima velocidade e economia de código.")

    add_h2("1.2 Histórico dos Padrões C e a Revolução do C99 para Embarcados")
    add_body("Na sequência do vídeo, o professor navega pelo slide de linha do tempo do C (B 1972, K&R 1978, C89/C90, C99, C11, C17). O professor destaca que a revisão C99 (1999) é a mais marcante para sistemas embarcados. Foi nela que surgiu a biblioteca <stdint.h> com os inteiros de largura fixa (uint8_t, uint16_t, uint32_t), além de comentários de linha dupla (//) e inicializações no meio do bloco.")

    add_h2("1.3 Esclarecimento: Existe uma linguagem 'Embedded C'?")
    add_body("O professor faz uma observação importante sobre o termo 'Embedded C': tecnicamente, não existe uma linguagem nova ou um compilador separado. O termo se refere à extensão ISO/IEC TR 18037 que padronizou o suporte do C para aritmética de ponto fixo, gerenciamento de espaços nomeados de memória e acesso direto aos registradores de hardware.")

    add_h2("1.4 Tabela de Palavras Reservadas (Keywords)")
    add_body("A aula apresenta a tabela com as 32 palavras reservadas da linguagem C (auto, break, case, char, const, continue, default, do, double, else, enum, extern, float, for, goto, if, inline, int, long, register, restrict, return, short, signed, sizeof, static, struct, switch, typedef, union, unsigned, void, volatile, while). O professor comenta rapidamente que no C11 foram adicionadas palavras especiais como _Alignas, _Atomic, _Generic, _Noreturn e _Static_assert.")

    add_h2("1.5 Organização de Arquivos e Cabeçalhos (APIs e Include Guards)")
    add_body("O professor explica que um arquivo de inclusão (.h) deve ser projetado como uma API, expondo apenas protótipos de funções públicas. Para evitar inclusão duplicada, ensina-se a usar as travas do pré-processador (#ifndef DEMO_H_ #define DEMO_H_ ... #endif). Ele comenta rapidamente também sobre o bloco extern \"C\", que previne a decoração de nomes (Name Mangling) quando o código C é vinculado a projetos em C++.")

    add_h2("1.6 O Especificador extern")
    add_body("O professor explica que a palavra-chave extern é usada para avisar ao compilador que uma variável global ou função citada naquele arquivo está declarada e alocada fisicamente em outro arquivo .c do projeto.")

    add_h2("1.7 Mapeamento de Memória (Flash ROM, SRAM e Seções do Linker Script)")
    add_body("A aula detalha como o compilador distribui o programa na memória física do STM32 através do arquivo de linkagem (.ld):")
    add_body("Armazena permanentemente o binário compilado e as variáveis const. Não apaga ao desligar a placa.", bold_prefix="• Seção .text (Flash ROM - 0x08000000): ")
    add_body("Guarda as variáveis globais e estáticas inicializadas. O boot copia os valores da Flash para a RAM.", bold_prefix="• Seção .data (SRAM - 0x20000000): ")
    add_body("Guarda as variáveis globais não inicializadas. O boot zera toda essa área.", bold_prefix="• Seção .bss (SRAM - 0x20000000): ")
    add_body("Área usada para alocar variáveis locais e registradores temporários de funções. O professor faz um alerta rigoroso sobre o perigo de estouro da pilha (Stack Overflow) se abusarmos de vetores locais grandes ou chamadas de funções recursivas.", bold_prefix="• Pilha (Stack): ")
    add_body("Usada para alocação dinâmica (malloc). O professor comenta que no desenvolvimento Bare-Metal o uso de malloc é desaconselhado por causar fragmentação imprevisível.", bold_prefix="• Monte (Heap): ")

    add_figure(r'C:\Users\vicen\Downloads\Video\frames_v1\frame_61m59s.png', "Figura 2: Slide do minuto 61:59 detalhando as seções de memória (.text, .data, .bss) e o arquivo de linkagem (.ld).")

    add_h2("1.8 Especificadores de Armazenamento: volatile, const, static, auto e register")
    add_body("O professor dá um foco enorme para a palavra-chave volatile. Ela avisa ao compilador que a variável pode mudar a qualquer momento por hardware externo ou interrupção (ISR), desativando otimizações que travariam o valor nos registradores da CPU (R0..R12).")
    add_body("O qualificador const força o armazenamento de tabelas na Flash ROM. O static preserva o valor de variáveis locais entre chamadas de função e restringe a visibilidade de variáveis globais apenas ao próprio arquivo .c. O professor comenta rapidamente sobre o especificador register (que solicita salvar a variável direto nos registradores internos da CPU) e auto (escopo padrão local).")

    add_h2("1.9 Manipulação de Registradores com Bitwise")
    add_body("Na etapa final do vídeo 1, o professor demonstra como alterar bits individuais em registradores de 32 bits através de operações bitwise sem destruir as configurações dos outros pinos: usa-se OR (|) para SET (ligar pino), AND com NOT (& ~) para CLEAR (desligar pino), XOR (^) para TOGGLE (inverter pino), e MASK com AND (&) para testar se um pino de entrada está em nível alto.")

    # --------------------------------------------------------------------------
    # VÍDEO 2: ABSTRAÇÃO E ESTRUTURAS AVANÇADAS
    # --------------------------------------------------------------------------
    add_h1("2. Acompanhamento Detalhado do Vídeo 2: Embedded C - Parte 2/2")

    add_h2("2.1 Passagem por Valor vs. Passagem por Referência")
    add_body("O segundo vídeo começa analisando a passagem de parâmetros em funções: na passagem por valor, os dados são copiados para a pilha (Stack); na passagem por referência, a função recebe o ponteiro do dado, economizando memória na pilha e permitindo alterar a variável original.")

    add_h2("2.2 Ponteiros, Arrays e Aritmética de Ponteiros")
    add_body("O professor explica que em C, o nome de um vetor é um ponteiro constante para o seu primeiro elemento. Ao somar +1 a um ponteiro, a CPU avança o tamanho exato em bytes do tipo de dado apontado. Ele cita rapidamente o tipo uintptr_t da <stdint.h>, usado para converter ponteiros em números inteiros do tamanho da palavra da arquitetura sem gerar alertas do compilador.")

    add_h2("2.3 Ponteiros para Função, Tabela IVT e Callbacks da HAL")
    add_body("O professor explica que o nome de qualquer função é um ponteiro para o seu endereço de início na Flash. Ele conecta isso diretamente com a Tabela de Vetores de Interrupção (IVT alocada em 0x08000000): quando ocorre uma interrupção de hardware (como o estouro do Timer TIM2), a CPU lê o ponteiro da tabela e salta direto para a função ISR de callback (ex.: HAL_TIM_PeriodElapsedCallback).")

    add_figure(r'C:\Users\vicen\Downloads\Video\frames_v2\frame_38m59s.png', "Figura 3: Slide do minuto 38:59 apresentando a sintaxe e a aplicação de Ponteiros para Função.")

    add_h2("2.4 Demonstração Prática no Godbolt Compiler Explorer")
    add_body("No minuto 62:59 do vídeo 2, o professor abre ao vivo o site Godbolt Compiler Explorer (godbolt.org) com o compilador ARM GCC 11.2 (none). Ele digita códigos C com volatile e const para mostrar no lado direito o código Assembly ARM gerado (instruções LDR/STR), demonstrando na prática como o compilador obedece a essas instruções.")

    add_figure(r'C:\Users\vicen\Downloads\Video\frames_v2\frame_62m59s.png', "Figura 4: Demonstração ao vivo no Godbolt Compiler Explorer (minuto 62:59) analisando as instruções Assembly geradas pelo ARM GCC.")

    add_h2("2.5 Qualificadores em Ponteiros de Leitura (const uint8_t *buf)")
    add_body("O professor recomenda a boa prática de declarar ponteiros de leitura de buffers com const (ex.: const uint8_t *buf em drivers spi_write ou lcd_write), garantindo que a função não altere o conteúdo original dos dados.")

    add_h2("2.6 Mapeamento de Registradores com struct (Memory-Mapped I/O)")
    add_body("O mapa de memória do ARM Cortex-M coloca os registradores de um periférico (como a porta GPIOA) em posições contíguas sequenciais (offsets 0x00, 0x04, 0x08, 0x0C). Ao criar uma struct que espelha essa sequência, podemos apontar um ponteiro para o endereço base da porta (0x40010800) e acessar os pinos com a sintaxe GPIOA->ODR, exatamente como é feito nas bibliotecas ST HAL e CMSIS.")

    add_figure(r'C:\Users\vicen\Downloads\Video\frames_v2\frame_70m59s.png', "Figura 5: Slide do minuto 70:59 demonstrando a organização de dados e periféricos através de estruturas (struct).")

    add_h2("2.7 Compartilhamento de Memória RAM com union")
    add_body("Diferente da struct, na union todos os membros compartilham a mesma posição física na memória RAM. O professor mostra que isso é muito útil para fatiar um dado numérico de 32 bits (uint32_t) em 4 bytes separados (uint8_t) para envio por pacotes USB HID ou comunicação serial UART sem precisar de máscaras manuais de bitwise.")

    add_figure(r'C:\Users\vicen\Downloads\Video\frames_v2\frame_78m59s.png', "Figura 6: Slide do minuto 78:59 explicando o funcionamento de Unióes (union) e o compartilhamento de RAM.")

    add_h2("2.8 Máquinas de Estados Finitos (FSM) com enum")
    add_body("Substituir números inteiros soltos por nomes de estados legíveis (enum) permite construir Máquinas de Estados Finitos (FSM) conectadas a laços switch-case, aumentando a segurança do código e evitando que o sistema caia em estados inválidos.")

    add_figure(r'C:\Users\vicen\Downloads\Video\frames_v2\frame_74m59s.png', "Figura 7: Slide do minuto 74:59 demonstrando a organização de enumerações (enum) para controle de fluxo.")

    add_h2("2.9 Exemplo Completo de Objeto Polimórfico (obj_t)")
    add_body("O professor apresenta um exemplo avançado reunindo struct, union e enum (obj_t) para processar diferentes tipos de dados e formatos geométricos no mesmo buffer de memória sem desperdiçar RAM.")

    add_h2("2.10 Organização do Modelo de Código-Fonte do Projeto")
    add_body("A aula se encerra exibindo a estrutura de organização dos arquivos de um projeto C profissional: separação entre cabeçalhos de inclusão de linguagem, inclusões de projetos, macros #define, protótipos de função e variáveis externas.")

    # --------------------------------------------------------------------------
    # SÍNTESE COMPARATIVA FINAL
    # --------------------------------------------------------------------------
    add_h1("3. Conclusão e Síntese Comparativa")
    add_body("Abaixo apresento a tabela resumo comparando as principais diferenças observadas entre a programação em C para computadores e para sistemas embarcados:")

    table_comp = doc.add_table(rows=1, cols=3)
    table_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table_comp.rows[0].cells
    hdr_cells[0].text = "Critério Técnico"
    hdr_cells[1].text = "C para PC (Desktop)"
    hdr_cells[2].text = "Embedded C (Embarcados)"

    for c in hdr_cells:
        set_cell_background(c, "232679")
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = 'Calibri'
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

    data = [
        ("Ambiente de Execução", "Aplicativo sobre Sistema Operacional", "Firmware Bare-Metal no silício"),
        ("Acesso ao Hardware", "Bloqueado pelo SO / Via Syscalls", "Direto por registradores MMIO"),
        ("Ciclo de Vida (main)", "Retorna 0 para o SO ao terminar", "Loop infinito while(1) ou repouso __WFI()"),
        ("Determinismo de Tempo", "Não-determinístico (escalonado pelo SO)", "Tempo Real determinístico por ISR/NVIC"),
        ("Qualificador volatile", "Raramente necessário", "Obrigatório em ISRs e registradores"),
        ("Tipos de Dados Inteiros", "Uso de int de tamanho variável", "Obrigatório <stdint.h> (uint8_t, uint32_t)"),
        ("Gerenciamento de RAM", "Abundante com memória virtual", "Estrito e contado (SRAM vs Flash ROM)")
    ]

    for row_data in data:
        row = table_comp.add_row()
        for idx, text in enumerate(row_data):
            cell = row.cells[idx]
            cell.text = text
            if idx == 0:
                set_cell_background(cell, "F1F5F9")
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx != 0 else WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = 'Calibri'
                    r.font.size = Pt(9.5)
                    if idx == 0:
                        r.bold = True

    # SALVAR ARQUIVO DOCX
    out_docx1 = r'C:\Users\vicen\Downloads\Vicenzo_Olivalves_Relatorio_Embedded_C_UFU.docx'
    out_docx2 = r'c:\Users\vicen\Downloads\streamdeck\Vicenzo_Olivalves_Relatorio_Embedded_C_UFU.docx'
    doc.save(out_docx1)
    doc.save(out_docx2)
    print(f"Successfully created EXHAUSTIVE Word Report V2 with Floating Point section: {out_docx1}")

if __name__ == '__main__':
    create_exhaustive_word_document_v2()
